import json
import os
import math
import pandas as pd
import geopandas as gpd

# Config paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MANUAL_DIR = os.path.join(BASE_DIR, 'manual_2026')

def load_json(filename):
    with open(os.path.join(MANUAL_DIR, filename), 'r', encoding='utf-8') as f:
        return json.load(f)

# Load configurations
TABLA_COBERTURAS = load_json('tabla3_coberturas.json')
TABLA_AMENAZAS = load_json('tabla4_amenazas.json')
TABLA_RANGOS = load_json('tabla_rangos.json')
FACTORES_EFECTIVIDAD = load_json('factores_efectividad.json')

def calcular_criterio_a(cobertura):
    """Devuelve (valor_A, es_natural). Si la cobertura no existe, asume transformada (1.0)."""
    if cobertura in TABLA_COBERTURAS:
        info = TABLA_COBERTURAS[cobertura]
        return info['valor'], info['tipo'] == 'natural'
    return 1.0, False

def calcular_criterio_b(df_arboles):
    """
    B = Σ(individuos_categoría × valor_categoría) / total_individuos
    df_arboles debe tener columna 'amenaza'. Si es nulo, se asume 'LC' (0.0).
    """
    if df_arboles.empty:
        return 0.0
    
    amenazas = df_arboles['amenaza'].fillna('LC').str.upper()
    # Mapear a LC si el string no está en la tabla
    valores = amenazas.map(lambda x: TABLA_AMENAZAS.get(x, 0.0))
    return valores.sum() / len(df_arboles)

def calcular_criterio_c(df_arboles):
    """
    CM = S / N (especies / individuos)
    Retorna el valor C discretizado por 0.1
    """
    if df_arboles.empty:
        return 0.1
    
    s = df_arboles['especie'].nunique()
    n = len(df_arboles)
    cm = s / n
    
    # Discretizar en pasos de 0.1 (ej. 0.0-0.1 -> 0.1; 0.11-0.20 -> 0.2)
    valor_c = math.ceil(cm * 10) / 10
    
    # Clamp min 0.1, max 1.0
    return max(0.1, min(1.0, valor_c))

def procesar_inventario(df_inventario):
    """
    Recibe el DataFrame del inventario (con columnas: id_arbol, especie, dap_cm, altura_total_m, cobertura, amenaza)
    Calcula el FCAFU por cada cobertura.
    Retorna un diccionario: { "Nombre Cobertura": { "tipo": "natural", "fcafu": 2.5, "A": 0.8, "B": 0.1, "C": 0.2, "count": 50 } }
    """
    resultados = {}
    
    # Validación de columnas
    columnas_req = ['especie', 'cobertura']
    for col in columnas_req:
        if col not in df_inventario.columns:
            raise ValueError(f"El inventario no tiene la columna obligatoria '{col}'")
            
    # Agrupar por cobertura
    for cobertura, group in df_inventario.groupby('cobertura'):
        val_A, es_natural = calcular_criterio_a(cobertura)
        
        if es_natural:
            val_B = calcular_criterio_b(group)
            val_C = calcular_criterio_c(group)
            fcafu_bruto = val_A + val_B + val_C
            fcafu_final = max(1.6, min(4.0, fcafu_bruto))  # Clamp manual 2026
            
            resultados[cobertura] = {
                "tipo": "natural",
                "fcafu": fcafu_final,
                "A": val_A,
                "B": val_B,
                "C": val_C,
                "count": len(group)
            }
        else:
            # Cobertura transformada, no aplica criterios A,B,C. Compensación 1:1
            resultados[cobertura] = {
                "tipo": "transformada",
                "fcafu": 1.0,
                "A": None,
                "B": None,
                "C": None,
                "count": len(group)
            }
            
    return resultados

def calcular_area_a_compensar(fcafu, factor_rango, area_impacto_ha):
    """ AC = (FCAFU + factor_rango) * Ai """
    return (fcafu + factor_rango) * area_impacto_ha

def calcular_adicionalidad(area_compensar_ha, accion):
    """
    Retorna las hectáreas adicionales esperadas según la acción (Conservar o Restaurar).
    """
    if accion == 'Conservar':
        tasa_bau = FACTORES_EFECTIVIDAD['Tasa_BAU']
        horizonte = FACTORES_EFECTIVIDAD['Horizonte_Anios']
        factor = FACTORES_EFECTIVIDAD['Conservar']
        return area_compensar_ha * tasa_bau * horizonte * factor
    elif accion == 'Restaurar':
        factor = FACTORES_EFECTIVIDAD['Restaurar']
        return area_compensar_ha * factor
    else:
        return 0.0

# ---------------------------------------------------------
# FASE 2: LÓGICA ESPACIAL CON GOOGLE EARTH ENGINE
# ---------------------------------------------------------
import ee
import geemap

def init_gee():
    """Inicializa la API de Earth Engine autenticando con el proyecto correcto."""
    try:
        ee.Initialize(project='ndvi-restauracion')
    except Exception as e:
        print("Error inicializando GEE, por favor autenticar con 'earthengine authenticate'")
        raise e

# Coberturas excluidas según la configuración JS
COBERTURAS_EXCLUIDAS = [
    'Tejido urbano continuo', 'Tejido urbano discontinuo', 'Zonas industriales o comerciales',
    'Red vial, ferroviaria y terrenos asociados', 'Aeropuertos', 'Lagunas, lagos y ciénagas naturales',
    'Ríos (50 m)', 'Cuerpos de agua artificiales', 'Cultivos transitorios', 'Cultivos permanentes herbáceos',
    'Cultivos permanentes arbustivos', 'Cultivos permanentes arbóreos', 'Cultivos agroforestales',
    'Cultivos confinados'
]

def obtener_contexto_geografico(impacto_geom_ee):
    """
    Recibe la geometría del impacto (ee.Geometry) y retorna el contexto.
    Retorna: (BIOMA_IMPACTO, NOM_ZH, NOM_SZH, NOM_MUNICIPIO)
    """
    ecosistemas = ee.FeatureCollection('projects/ndvi-restauracion/assets/Shape_E_ECCMC_Ver21_100K')
    zh_colombia = ee.FeatureCollection('projects/ndvi-restauracion/assets/zh_colombia')
    municipios = ee.FeatureCollection('FAO/GAUL/2015/level2').filter(ee.Filter.eq('ADM0_NAME', 'Colombia'))
    
    eco_impactado = ecosistemas.filterBounds(impacto_geom_ee).first()
    bioma = ee.String(eco_impactado.get('BIOMA_IAvH')).getInfo()
    
    zh_impactada = zh_colombia.filterBounds(impacto_geom_ee).first()
    nom_zh = ee.String(zh_impactada.get('nom_zh')).getInfo()
    nom_szh = ee.String(zh_impactada.get('nom_szh')).getInfo()
    
    mun_impactado = municipios.filterBounds(impacto_geom_ee).first()
    nom_mun = ee.String(mun_impactado.get('ADM2_NAME')).getInfo()
    
    return bioma, nom_zh, nom_szh, nom_mun

def construir_rango_gee(filtro_bioma, geometria_zona, nombre_rango, factor_adicional, restar_runap):
    """
    Construye las geometrías de un rango.
    Retorna un diccionario con las FeatureCollections resultantes en Earth Engine.
    """
    ecosistemas = ee.FeatureCollection('projects/ndvi-restauracion/assets/Shape_E_ECCMC_Ver21_100K')
    runap = ee.FeatureCollection('projects/ndvi-restauracion/assets/RUNAP')
    
    candidatos = ecosistemas.filter(filtro_bioma).filterBounds(geometria_zona)
    
    # Separar excluidos
    excluidos_cobertura = candidatos.filter(ee.Filter.inList('COBERTURA', COBERTURAS_EXCLUIDAS))
    candidatos_validos = candidatos.filter(ee.Filter.inList('COBERTURA', COBERTURAS_EXCLUIDAS).Not())
    
    runap_en_zona = runap.filterBounds(geometria_zona)
    exclusiones_geom = runap_en_zona.geometry() if restar_runap else ee.Geometry.MultiPolygon([])
    
    # Mapeo de recorte
    def recortar_y_etiquetar(f):
        geom_zona = f.geometry().intersection(geometria_zona, 1)
        geom_neta = geom_zona.difference(exclusiones_geom, 1)
        area_ha = geom_neta.area().divide(10000)
        return f.setGeometry(geom_neta).set('area_ha', area_ha).set('rango', nombre_rango).set('fact_adic', factor_adicional)
        
    procesados = candidatos_validos.map(recortar_y_etiquetar)
    procesados = procesados.filter(ee.Filter.gt('area_ha', 0.01))
    
    conservar = procesados.filter(ee.Filter.eq('GRADO_TRAN', 'Natural')).map(lambda f: f.set('accion', 'Conservar'))
    restaurar = procesados.filter(ee.Filter.eq('GRADO_TRAN', 'Transformado')).map(lambda f: f.set('accion', 'Restaurar'))
    
    return {
        'nombre_rango': nombre_rango,
        'factor_adicional': factor_adicional,
        'conservar': conservar,
        'restaurar': restaurar,
        'excluidos_cobertura': excluidos_cobertura,
        'runap_en_zona': runap_en_zona,
        'geometria_zona': geometria_zona
    }

def generar_rangos(impacto_geom_ee):
    """
    Genera y retorna la configuración de los rangos R1 a R5.
    (R6 puede implementarse como excepción nacional).
    Retorna dict con datos de GEE listos para exportar a GeoDataFrame.
    """
    bioma, nom_zh, nom_szh, nom_mun = obtener_contexto_geografico(impacto_geom_ee)
    
    municipios = ee.FeatureCollection('FAO/GAUL/2015/level2').filter(ee.Filter.eq('ADM0_NAME', 'Colombia'))
    zh_colombia = ee.FeatureCollection('projects/ndvi-restauracion/assets/zh_colombia')
    
    area_influencia = municipios.filter(ee.Filter.eq('ADM2_NAME', nom_mun)).geometry()
    poligono_szh = zh_colombia.filter(ee.Filter.eq('nom_szh', nom_szh)).geometry()
    poligono_zh = zh_colombia.filter(ee.Filter.eq('nom_zh', nom_zh)).geometry()
    
    mismo_bioma = ee.Filter.eq('BIOMA_IAvH', bioma)
    otro_bioma = ee.Filter.neq('BIOMA_IAvH', bioma)
    
    rangos = {
        'R1': construir_rango_gee(mismo_bioma, area_influencia, 'R1', TABLA_RANGOS['R1'], True),
        'R2': construir_rango_gee(mismo_bioma, poligono_szh, 'R2', TABLA_RANGOS['R2'], True),
        # R3, R4, R5 se implementan según se requiera para evitar timeouts
    }
    
    contexto = {
        'BIOMA_IAvH': bioma,
        'ZH': nom_zh,
        'SZH': nom_szh,
        'Municipio': nom_mun
    }
    
    return contexto, rangos

def descargar_rango_gdf(rango_dict):
    """
    Recibe el diccionario de un rango y descarga 'conservar' y 'restaurar' a un GeoDataFrame.
    Realiza un fallback con try/except si es muy grande.
    """
    conservar_fc = rango_dict['conservar']
    restaurar_fc = rango_dict['restaurar']
    
    # Combinar ambas para descargar una sola vez
    combinado = conservar_fc.merge(restaurar_fc)
    
    try:
        gdf = geemap.ee_to_gdf(combinado)
        return gdf
    except Exception as e:
        print(f"Error descargando {rango_dict['nombre_rango']}: {e}. Fallback necesario a Export.")
        return None

# ---------------------------------------------------------
# FASE 4: EXPORTACIÓN (DOCX y ZIP)
# ---------------------------------------------------------
import io
import zipfile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_kmz(uploaded_file):
    """
    Lee un archivo KMZ subido a Streamlit, lo convierte a GeoDataFrame
    y retorna el gdf y el área de impacto en hectáreas.
    """
    import fiona
    fiona.drvsupport.supported_drivers['KML'] = 'rw'
    fiona.drvsupport.supported_drivers['LIBKML'] = 'rw'
    
    with open("temp.kmz", "wb") as f:
        f.write(uploaded_file.getvalue())
        
    gdf = gpd.read_file("temp.kmz", driver="KML")
    # Reproject to Web Mercator to get area in square meters
    gdf_proj = gdf.to_crs(epsg=3857)
    area_ha = gdf_proj.geometry.area.sum() / 10000
    
    # Clean up temp file
    if os.path.exists("temp.kmz"):
        os.remove("temp.kmz")
        
    return gdf, area_ha

def gdf_to_ee_poly(gdf):
    """
    Convierte el primer polígono del GeoDataFrame a ee.Geometry
    """
    geom = gdf.geometry.iloc[0]
    # Si es multi-polígono, podemos simplificar o agarrar la envolvente
    # para el MVP, asumiremos un solo polígono o usaremos coord list.
    if geom.geom_type == 'Polygon':
        coords = list(geom.exterior.coords)
        return ee.Geometry.Polygon(coords)
    elif geom.geom_type == 'MultiPolygon':
        polys = []
        for poly in geom.geoms:
            polys.append(list(poly.exterior.coords))
        return ee.Geometry.MultiPolygon(polys)
    return None

def generar_reporte_docx(contexto, resultados_fcafu, area_impacto, rangos_ac, rangos_adic):
    """
    Genera el reporte DOCX en memoria.
    """
    doc = Document()
    
    # Título Principal
    titulo = doc.add_heading('Plan de Compensación del Componente Biótico', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('1. Datos del Proyecto', level=1)
    doc.add_paragraph(f"Municipio: {contexto['Municipio']}")
    doc.add_paragraph(f"BIOMA-IAvH: {contexto['BIOMA_IAvH']}")
    doc.add_paragraph(f"Zona Hidrográfica (ZH): {contexto['ZH']}")
    doc.add_paragraph(f"Subzona Hidrográfica (SZH): {contexto['SZH']}")
    doc.add_paragraph(f"Área de Impacto (Ai): {area_impacto:.2f} hectáreas")
    
    doc.add_heading('2. Memoria de Cálculo del FCAFU', level=1)
    for cobertura, datos in resultados_fcafu.items():
        doc.add_heading(f"Cobertura: {cobertura}", level=2)
        if datos['tipo'] == 'natural':
            doc.add_paragraph(f"Criterio A (Tabla 3): {datos['A']}")
            doc.add_paragraph(f"Criterio B (Tabla 4, datos de inventario): {datos['B']:.3f}")
            doc.add_paragraph(f"Criterio C (Tabla 5): {datos['C']:.3f}")
            doc.add_paragraph(f"FCAFU Total Calculado: {datos['fcafu']:.3f}")
        else:
            doc.add_paragraph("Cobertura Transformada - Compensación 1:1 según Manual 2026.")
            doc.add_paragraph(f"FCAFU Aplicado: {datos['fcafu']:.1f}")

    doc.add_heading('3. Áreas Candidatas por Rango y Área a Compensar (AC)', level=1)
    # Ejemplo de inserción de tabla
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rango'
    hdr_cells[1].text = 'Factor Adicional'
    hdr_cells[2].text = 'Conservar (ha)'
    hdr_cells[3].text = 'Restaurar (ha)'
    
    for r_name, ac_data in rangos_ac.items():
        row_cells = table.add_row().cells
        row_cells[0].text = r_name
        row_cells[1].text = f"+{TABLA_RANGOS.get(r_name, 0.0)}"
        row_cells[2].text = f"{ac_data.get('conservar', 0):.2f}"
        row_cells[3].text = f"{ac_data.get('restaurar', 0):.2f}"

    doc.add_heading('4. Adicionalidad Esperada', level=1)
    doc.add_paragraph("Basado en los siguientes factores de efectividad comprobados:")
    doc.add_paragraph("Conservar: 0.85 (Andam et al. 2008, Pfaff et al. 2014)")
    doc.add_paragraph("Restaurar: 0.75 (Crouzeilles et al. 2017, González-M. et al. 2018)")
    
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr2_cells = table2.rows[0].cells
    hdr2_cells[0].text = 'Rango'
    hdr2_cells[1].text = 'AC Requerida Total (ha)'
    hdr2_cells[2].text = 'Adic. 100% Conservar (ha)'
    hdr2_cells[3].text = 'Adic. 100% Restaurar (ha)'
    
    for r_name, adic_data in rangos_adic.items():
        row_cells = table2.add_row().cells
        row_cells[0].text = r_name
        row_cells[1].text = f"{adic_data.get('ac_total', 0):.2f}"
        row_cells[2].text = f"{adic_data.get('adic_cons', 0):.2f}"
        row_cells[3].text = f"{adic_data.get('adic_rest', 0):.2f}"

    doc.add_heading('5. Bibliografía', level=1)
    doc.add_paragraph("- Andam et al. (2008) PNAS. DOI: 10.1073/pnas.0800437105")
    doc.add_paragraph("- Pfaff et al. (2014) World Dev. DOI: 10.1016/j.worlddev.2013.01.011")
    doc.add_paragraph("- Crouzeilles et al. (2017) Sci Adv. DOI: 10.1126/sciadv.1701345")
    doc.add_paragraph("- González-M. et al. (2018) IAvH. Catálogo BST Colombia")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_zip_descargables(gdfs_rangos, excel_buffer):
    """
    Genera un archivo ZIP con los shapefiles y el Excel.
    """
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        # Añadir Excel
        zip_file.writestr("memoria_calculo.xlsx", excel_buffer.getvalue())
        
        # Añadir GDFs como GeoJSON para no crear multiles archivos .shp .dbf .shx
        for r_name, gdf in gdfs_rangos.items():
            if gdf is not None and not gdf.empty:
                geojson_str = gdf.to_json()
                zip_file.writestr(f"{r_name}_candidatas.geojson", geojson_str)
                
    zip_buffer.seek(0)
    return zip_buffer

if __name__ == '__main__':
    print("Core metrics module loaded successfully.")

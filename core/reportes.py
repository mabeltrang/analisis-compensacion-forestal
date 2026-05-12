import pandas as pd
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import settings

def generar_reporte_excel(datos, output_path):
    """Genera el Excel de resultados con mltiples hojas segn el Manual 2026"""
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        # 1. Hoja Resumen
        resumen_data = {
            'Variable': ['Proyecto', 'Cdigo', 'Municipio', 'Departamento', 'Bioma Impacto', 'Zona Hidrogrfica', 'Tasa BAU Bioma'],
            'Valor': [
                datos['proyecto'], 
                datos['codigo'], 
                datos['contexto']['municipio'], 
                datos['contexto']['departamento'],
                datos['contexto']['bioma_principal'],
                datos['contexto']['zh'],
                f"{datos['bau']['tasa_bau_anual']*100:.4f}%"
            ]
        }
        pd.DataFrame(resumen_data).to_excel(writer, sheet_name='Resumen', index=False)
        
        # 2. Hoja Comparativa de Rangos (TABLA PRINCIPAL)
        comparativa = []
        for rango, r_data in datos['atc'].items():
            candidata = datos['candidatas'].get(rango, {})
            comparativa.append({
                'Rango Geogrfico': rango,
                'Factor Adicional': r_data['factor_adicional'],
                'ATC Requerido (ha)': round(r_data['atc_total'], 2),
                'Hectreas Conservar': round(candidata.get('ha_conservar', 0), 2),
                'Hectreas Restaurar': round(candidata.get('ha_restaurar', 0), 2),
                'Total Disponible (ha)': round(candidata.get('total', 0), 2),
                'Cubre Requerimiento?': 'SI' if (candidata.get('total', 0) >= r_data['atc_total']) else 'NO',
                'Prdida BAU Evitada (ha/ao)': round(candidata.get('total', 0) * datos['bau']['tasa_bau_anual'], 4)
            })
        pd.DataFrame(comparativa).to_excel(writer, sheet_name='Comparativa_Rangos', index=False)
        
        # 3. Hoja FCAFU Detallado por Cobertura
        fcafu_det = []
        for cob, inv in datos['inventario_full'].items():
            fcafu_det.append({
                'Cobertura': cob,
                'Individuos (N)': inv['N'],
                'Especies (S)': inv['S'],
                'Mezcla (S/N)': round(inv['SN'], 3),
                'Criterio A (Ecosistema)': inv['A'],
                'Criterio B (Amenaza)': round(inv['B'], 3),
                'Criterio C (Composicin)': inv['C'],
                'FCAFU Resultante': round(inv['FCAFU'], 3),
                'rea Basal Total (m2)': round(inv['area_basal_total'], 3)
            })
        pd.DataFrame(fcafu_det).to_excel(writer, sheet_name='Detalle_FCAFU', index=False)
        
        # 4. Hoja Especies Amenazadas
        amenazadas_list = []
        for cob, inv in datos['inventario_full'].items():
            for sp in inv['amenazadas']:
                amenazadas_list.append({
                    'Cobertura': cob,
                    'Nombre Cientfico': sp['Nombre cientifico'],
                    'Categora': sp['categoria_amenaza']
                })
        if amenazadas_list:
            pd.DataFrame(amenazadas_list).to_excel(writer, sheet_name='Especies_Amenazadas', index=False)

def generar_reporte_word(datos, output_path):
    """Genera el informe tcnico siguiendo la estructura GF-PN-01 de Unergy"""
    doc = Document()
    
    # Estilo de Ttulo
    title = doc.add_heading(f"INFORME TCNICO DE COMPENSACIN FORESTAL", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Proyecto: {datos['proyecto']}").bold = True
    doc.add_paragraph(f"Cdigo Interno: {datos['codigo']}")
    doc.add_paragraph("---")
    
    # Seccin 1: Punto de Partida
    doc.add_heading("1. PUNTO DE PARTIDA Y CONTEXTO", level=1)
    p = doc.add_paragraph()
    p.add_run("El rea de impacto se localiza en el municipio de ").add_run(f"{datos['contexto']['municipio']}").bold = True
    p.add_run(f", departamento de {datos['contexto']['departamento']}. ")
    p.add_run(f"Pertenece al Bioma IAvH ").add_run(f"{datos['contexto']['bioma_principal']}").bold = True
    p.add_run(f" y a la Zona Hidrogrfica de {datos['contexto']['zh']}.")
    
    # Seccin 2: Clculo del FCAFU
    doc.add_heading("2. CLCULO DEL FACTOR DE COMPENSACIN (FCAFU)", level=1)
    doc.add_paragraph("De acuerdo al Manual 2026, el clculo se realiz por cada cobertura presente:")
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(['Cobertura', 'A', 'B', 'C', 'FCAFU', 'N']):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    for cob, inv in datos['inventario_full'].items():
        row = table.add_row().cells
        row[0].text = str(cob)
        row[1].text = f"{inv['A']:.2f}"
        row[2].text = f"{inv['B']:.3f}"
        row[3].text = f"{inv['C']:.2f}"
        row[4].text = f"{inv['FCAFU']:.3f}"
        row[5].text = str(inv['N'])
        
    # Seccin 3: Adicionalidad
    doc.add_heading("3. ANLISIS DE ADICIONALIDAD", level=1)
    doc.add_paragraph(f"La tasa BAU (Business As Usual) para el bioma afectado es de {datos['bau']['tasa_bau_anual']*100:.4f}% anual. ")
    doc.add_paragraph("Este anlisis integra:")
    bullets = [
        "Tasa de prdida anual de cobertura (Hansen GFC).",
        "Trayectoria sucesional de las coberturas naturales.",
        "Mapeo de iniciativas de conservacin existentes (SINAP).",
        "Escenarios comparativos entre los 5 rangos geogrficos."
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
        
    # Seccin 4: Seleccin Recomendada
    doc.add_heading("4. SELECCIN RECOMENDADA", level=1)
    r1_data = datos['atc']['Rango 1']
    r1_cand = datos['candidatas']['Rango 1']
    
    if r1_cand['total'] >= r1_data['atc_total']:
        doc.add_paragraph("Se recomienda priorizar el RANGO 1 (rea de Influencia), dado que existe disponibilidad de hectreas para cubrir el ATC requerido.")
    else:
        doc.add_paragraph("Se recomienda escalar al RANGO 2 o 3, debido a que el Rango 1 no cuenta con suficiencia de reas candidatas elegibles.")
        
    doc.save(output_path)
